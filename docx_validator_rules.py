from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
import re


class DocumentValidator:
    def __init__(self, doc_path):
        self.doc_path = doc_path
        self.doc = Document(doc_path)
        self.issues = []

    def validate_all(self):
        self.issues = []

        self.validate_page_format()
        self.validate_font()
        self.validate_paragraphs()
        self.validate_headings()
        self.validate_page_numbering()
        self.validate_tables()
        self.validate_figures()
        self.validate_equations()
        self.validate_references()
        self.validate_appendices()

        return self.issues

    def validate_page_format(self):
        try:
            from docx.shared import Mm
            
            a4_height = Mm(297)
            a4_width = Mm(210)
            left_margin_required = Mm(30)
            right_margin_required = Mm(15)
            right_margin_alternative = Mm(10)
            top_margin_required = Mm(20)
            bottom_margin_required = Mm(20)
            
            page_tolerance = Mm(1)
            margin_tolerance = Mm(1)

            for section in self.doc.sections:
                if not (abs(section.page_width - a4_width) <= page_tolerance and
                        abs(section.page_height - a4_height) <= page_tolerance):
                    width_mm = round(section.page_width.mm, 1)
                    height_mm = round(section.page_height.mm, 1)
                    self.issues.append(f"Размер страницы не соответствует формату A4 (210×297 мм). Текущий размер: {width_mm}×{height_mm} мм.")

                if abs(section.left_margin - left_margin_required) > margin_tolerance:
                    left_mm = round(section.left_margin.mm, 1)
                    self.issues.append(f"Левое поле должно быть 30 мм. Текущее: {left_mm} мм.")

                if (abs(section.right_margin - right_margin_required) > margin_tolerance and
                    abs(section.right_margin - right_margin_alternative) > margin_tolerance):
                    right_mm = round(section.right_margin.mm, 1)
                    self.issues.append(f"Правое поле должно быть 15 мм (допускается 10 мм). Текущее: {right_mm} мм.")

                if abs(section.top_margin - top_margin_required) > margin_tolerance:
                    top_mm = round(section.top_margin.mm, 1)
                    self.issues.append(f"Верхнее поле должно быть 20 мм. Текущее: {top_mm} мм.")

                if abs(section.bottom_margin - bottom_margin_required) > margin_tolerance:
                    bottom_mm = round(section.bottom_margin.mm, 1)
                    self.issues.append(f"Нижнее поле должно быть 20 мм. Текущее: {bottom_mm} мм.")

        except Exception as e:
            self.issues.append("Пожалуйста, проверьте формат A4 (210×297 мм) и поля (левое - 30 мм, правое - 15 мм, верхнее и нижнее - 20 мм) вручную.")

    def get_first_line_indent_cm(self, paragraph):
        from docx.oxml.ns import qn
        pPr = paragraph._p.pPr
        if pPr is None:
            return None
        ind = pPr.find(qn('w:ind'))
        if ind is None:
            return None
        val = ind.get(qn('w:firstLine'))
        if val is None:
            return None
        try:
            twips = float(val)
            cm = twips * 0.0017638889
            return round(cm, 3)
        except ValueError:
            return None

    def get_line_spacing_value(self, paragraph):
        from docx.oxml.ns import qn
        
        pPr = paragraph._p.pPr
        if pPr is not None:
            spacing = pPr.find(qn('w:spacing'))
            if spacing is not None:
                line_val = spacing.get(qn('w:line'))
                line_rule = spacing.get(qn('w:lineRule'))
                
                if line_val is not None:
                    try:
                        line_val = float(line_val)
                        
                        if line_rule is None:
                            line_rule = 'auto'
                            
                        if line_rule == 'auto':
                            line_spacing = line_val / 240.0
                            return line_spacing, line_rule
                        else:
                            points = line_val / 20.0
                            if points / 12.0 > 1.2:
                                line_spacing = points / 12.0
                            else:
                                line_spacing = 1.0
                            return line_spacing, line_rule
                    except (ValueError, TypeError):
                        pass
        
        try:
            line_spacing = paragraph.paragraph_format.line_spacing
            if line_spacing is not None:
                if 1.4 <= line_spacing <= 1.6:
                    return 1.5, 'auto'
                elif 1.9 <= line_spacing <= 2.1:
                    return 2.0, 'auto'
                elif 0.9 <= line_spacing <= 1.1:
                    return 1.0, 'auto'
                else:
                    return line_spacing, 'auto'
                    
            line_rule = paragraph.paragraph_format.line_spacing_rule
            if line_rule is not None:
                from docx.enum.text import WD_LINE_SPACING
                if line_rule == WD_LINE_SPACING.ONE_POINT_FIVE:
                    return 1.5, 'auto'
                elif line_rule == WD_LINE_SPACING.DOUBLE:
                    return 2.0, 'auto'
                elif line_rule == WD_LINE_SPACING.SINGLE:
                    return 1.0, 'auto'
        except:
            pass
            
        try:
            if paragraph.style:
                style_name = paragraph.style.name
                if "полуторный" in style_name.lower() or "1.5" in style_name:
                    return 1.5, 'auto'
                elif "двойной" in style_name.lower() or "2.0" in style_name or "double" in style_name.lower():
                    return 2.0, 'auto'
                elif "одинарный" in style_name.lower() or "1.0" in style_name or "single" in style_name.lower():
                    return 1.0, 'auto'
        except:
            pass
            
        return 1.0, 'auto'

    def validate_font(self):
        for paragraph in list(filter(lambda x: x != '', self.doc.paragraphs)):
            if not paragraph.text.strip():
                continue

            for run in paragraph.runs:
                if run.font.name and 'Times New Roman' not in run.font.name:
                    self.issues.append(f"Шрифт должен быть Times New Roman. Текущий: {run.font.name} в тексте: '{run.text[:20]}...'")

                if run.font.size is not None:
                    font_size_pt = run.font.size.pt
                    if font_size_pt < 12:
                        self.issues.append(f"Размер шрифта должен быть не менее 12 пт. Текущий: {font_size_pt} пт в тексте: '{run.text[:20]}...'")

    def validate_paragraphs(self):
        from docx.shared import Cm
        
        heading_patterns = [
            r'^(СОДЕРЖАНИЕ|ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|ПРИЛОЖЕНИЕ|СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ|СПИСОК СОКРАЩЕНИЙ И УСЛОВНЫХ ОБОЗНАЧЕНИЙ|ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ)',
            r'^(\d+(?:\.\d+)?(?:\.\d+)*)\s+[А-ЯЁ]',
            r'^Рисунок\s+\d+(?:\.\d+)?',
            r'^Таблица\s+\d+(?:\.\d+)?',
        ]
        
        in_terms_section = False
        in_abbreviations_section = False
        
        for paragraph in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
            text = paragraph.text.strip()
            if not text:
                continue
            
            if "ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ" in text.upper():
                in_terms_section = True
                in_abbreviations_section = False
            elif "СПИСОК СОКРАЩЕНИЙ И УСЛОВНЫХ ОБОЗНАЧЕНИЙ" in text.upper() or "СПИСОК СОКРАЩЕННЫХ И УСЛОВНЫХ ОБОЗНАЧЕНИЙ" in text.upper():
                in_abbreviations_section = True
                in_terms_section = False
            elif any(section in text.upper() for section in ["ВВЕДЕНИЕ", "СОДЕРЖАНИЕ", "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ", "ЗАКЛЮЧЕНИЕ", "ПРИЛОЖЕНИЕ"]):
                in_terms_section = False
                in_abbreviations_section = False
                
            is_heading = False
            for pattern in heading_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    is_heading = True
                    break
                    
            is_toc_entry = '\t' in paragraph.text or text.endswith('...')
            is_list_item = text.startswith('•') or text.startswith('-') or re.match(r'^\d+\.', text)
            is_table_continuation = re.match(r'^Продолжение таблицы \d+(?:\.\d+)?', text)
            
            is_justified = (paragraph.alignment == WD_ALIGN_PARAGRAPH.JUSTIFY or paragraph.alignment is None)
            
            if not is_heading and not is_toc_entry and not is_list_item and not is_table_continuation and not is_justified:
                self.issues.append(f"Основной текст должен быть выровнен по ширине. Текущее выравнивание: {paragraph.alignment} в тексте: '{text[:20]}...'")
            
            if is_heading or is_toc_entry or is_list_item or is_table_continuation:
                pass
            elif in_terms_section or in_abbreviations_section:
                pass
            elif re.match(r'^[A-ZА-Я]{2,}', text) or text.startswith('где '):
                pass
            else:
                current_indent = self.get_first_line_indent_cm(paragraph)
                if current_indent is None or (Cm(current_indent) < Cm(1.24) or Cm(current_indent) > Cm(1.26)):
                    self.issues.append(f"Абзацный отступ должен быть 1,25 см. Текущий: {current_indent if current_indent is not None else 0} см в тексте: '{text[:20]}...'")

                if not is_heading and not is_toc_entry and not is_list_item and not is_table_continuation:
                    try:
                        line_spacing, line_rule = self.get_line_spacing_value(paragraph)
                        
                        if not (1.4 <= line_spacing <= 1.6):
                            if line_rule == 'auto':
                                rule_text = "множитель"
                            else:
                                rule_text = "точно"
                                
                            self.issues.append(f"Межстрочный интервал должен быть полуторным (1.5). Текущий: {line_spacing:.2f} ({rule_text}) в тексте: '{text[:20]}...'")
                    except Exception as e:
                        pass

    def validate_headings(self):
        headings = []
        heading_patterns = [
            (r'^(СОДЕРЖАНИЕ|ВВЕДЕНИЕ|ЗАКЛЮЧЕНИЕ|ПРИЛОЖЕНИЕ|СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ|СПИСОК СОКРАЩЕНИЙ И УСЛОВНЫХ ОБОЗНАЧЕНИЙ|ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ)(\s*$|\.{3}.*)', 0),
            (r'^(\d+)\s+[А-ЯЁ]', 1),
            (r'^(\d+\.\d+)\s+[А-ЯЁ]', 2),
            (r'^(\d+\.\d+\.\d+)\s+[А-ЯЁ]', 3)
        ]

        for paragraph in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
            text = paragraph.text.strip()
            if not text:
                continue

            for pattern, level in heading_patterns:
                if re.match(pattern, text, re.IGNORECASE):
                    headings.append((text, level, paragraph))
                    break

        for text, level, paragraph in headings:
            if level == 0:
                if paragraph.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
                    if '\t' not in text and not text.endswith('...'):
                        self.issues.append(f"Заголовок '{text}' должен быть выровнен по центру.")

                if not text.isupper() and '\t' not in text and not text.endswith('...'):
                    if not text.split()[0].isupper():
                        self.issues.append(f"Заголовок '{text}' должен быть написан прописными буквами.")

            else:
                if text.endswith('.'):
                    self.issues.append(f"Заголовок '{text}' не должен оканчиваться точкой.")

                if level == 1:
                    try:
                        match = re.match(r'(\d+(?:\.\d+)?)', text)
                        if match:
                            number = int(float(match.group(1)))
                            if number < 1:
                                self.issues.append(f"Некорректная нумерация заголовка '{text}'.")
                    except (ValueError, TypeError):
                        self.issues.append(f"Некорректная нумерация заголовка '{text}'.")

    def validate_page_numbering(self):
        try:
            has_page_numbers = False
            page_numbers_start_on_third_page = False

            for section_index, section in enumerate(self.doc.sections):
                if hasattr(section, 'footer') and section.footer and section.footer.paragraphs:
                    footer_paragraphs = section.footer.paragraphs

                    if any(p.text.strip() for p in footer_paragraphs):
                        has_page_numbers = True
                        for p in footer_paragraphs:
                            if p.text.strip() and p.alignment not in (WD_ALIGN_PARAGRAPH.CENTER, None):
                                self.issues.append("Номера страниц должны быть размещены в центре нижней части страницы.")
                        
                        if section_index >= 2:
                            page_numbers_start_on_third_page = True
                        break

                    for p in footer_paragraphs:
                        if hasattr(p, '_p') and p._p is not None:
                            xml_str = p._p.xml
                            if 'PAGE' in xml_str or 'w:fldChar' in xml_str:
                                has_page_numbers = True
                                
                                if section_index >= 2:
                                    page_numbers_start_on_third_page = True
                                break

            if not has_page_numbers:
                for i, paragraph in enumerate(self.doc.paragraphs):
                    approx_page = (i // 20) + 1
                    
                    if "PAGE" in paragraph.text or "NUMPAGES" in paragraph.text:
                        has_page_numbers = True
                        if approx_page >= 3:
                            page_numbers_start_on_third_page = True
                        break

                    if hasattr(paragraph, '_p') and paragraph._p is not None:
                        xml_str = paragraph._p.xml
                        if 'PAGE' in xml_str or 'w:fldChar' in xml_str:
                            has_page_numbers = True
                            if approx_page >= 3:
                                page_numbers_start_on_third_page = True
                            break
            
            if not has_page_numbers and len(list(self.doc.paragraphs)) > 20:
                self.issues.append("Ошибка в нумерации страниц. Номера должны быть в центре нижней части страницы.")
            
            if has_page_numbers and not page_numbers_start_on_third_page:
                self.issues.append("Нумерация страниц должна начинаться с третьей страницы.")

        except Exception as e:
            pass

    def validate_tables(self):
        if not self.doc.tables:
            return

        table_numbers = {}
        table_captions = []

        for paragraph in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
            text = paragraph.text.strip()
            if text.startswith("Таблица "):
                match = re.match(r"^Таблица\s+(\d+(?:\.\d+)?)", text)
                if match:
                    try:
                        table_num = int(float(match.group(1)))
                        if table_num not in table_numbers:
                            table_numbers[table_num] = []
                        table_numbers[table_num].append((paragraph, False))
                        table_captions.append((table_num, paragraph))
                    except (ValueError, TypeError):
                        self.issues.append(f"Некорректный номер таблицы: '{match.group(1)}' в тексте: '{text[:30]}...'")
            elif text.startswith("Продолжение таблицы "):
                match = re.match(r"^Продолжение таблицы\s+(\d+(?:\.\d+)?)", text)
                if match:
                    try:
                        table_num = int(float(match.group(1)))
                        if table_num not in table_numbers:
                            table_numbers[table_num] = []
                        table_numbers[table_num].append((paragraph, True))
                        table_captions.append((table_num, paragraph))
                    except (ValueError, TypeError):
                        self.issues.append(f"Некорректный номер продолжения таблицы: '{match.group(1)}' в тексте: '{text[:30]}...'")

        tables_in_order = sorted(table_numbers.keys())

        for i, table_num in enumerate(tables_in_order, 1):
            if table_num != i:
                self.issues.append(f"Нарушена последовательность нумерации таблиц. Таблица с номером {table_num} должна иметь номер {i}.")

            captions = table_numbers[table_num]
            main_caption = None
            for cap, is_continuation in captions:
                if not is_continuation:
                    main_caption = cap
                    break

            if main_caption:
                if not re.match(r"^Таблица\s+" + str(table_num) + r"(\s+[-–—]\s+|\s*$)", main_caption.text.strip()):
                    self.issues.append(f"Таблица #{table_num} не имеет правильной подписи. Подпись должна быть в формате 'Таблица {table_num} - Наименование'.")

            for caption_paragraph, is_continuation in captions:
                if is_continuation:
                    if caption_paragraph.alignment != WD_ALIGN_PARAGRAPH.RIGHT:
                        self.issues.append(f"Продолжение таблицы '{caption_paragraph.text[:30]}...' должно быть выровнено по правому краю.")
                else:
                    if caption_paragraph.alignment != WD_ALIGN_PARAGRAPH.LEFT and caption_paragraph.alignment is not None:
                        self.issues.append(f"Подпись таблицы '{caption_paragraph.text[:30]}...' должна быть выровнена по левому краю.")

            has_reference = False
            ref_pattern = r'таблиц[аеуыи]?\s+' + str(table_num) + r'\b'

            for paragraph in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
                if re.search(ref_pattern, paragraph.text, re.IGNORECASE):
                    has_reference = True
                    break

            if not has_reference:
                self.issues.append(f"На таблицу #{table_num} нет ссылки в тексте. На все таблицы должны быть ссылки.")

    def validate_figures(self):
        figure_captions = []
        caption_paragraphs = []

        for paragraph in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
            if re.match(r'^Рисунок\s+\d+', paragraph.text.strip()):
                figure_captions.append(paragraph.text.strip())
                caption_paragraphs.append(paragraph)

        if not figure_captions:
            return

        for i, (caption, paragraph) in enumerate(zip(figure_captions, caption_paragraphs), 1):
            if not re.match(r'^Рисунок\s+\d+\s+[\-–—]\s+\w', caption):
                self.issues.append(f"Неправильный формат подписи рисунка: '{caption}'. Должно быть 'Рисунок Номер - Наименование'.")

            if paragraph.alignment != WD_ALIGN_PARAGRAPH.CENTER and paragraph.alignment is not None:
                self.issues.append(f"Подпись рисунка '{caption[:30]}...' должна быть выровнена по центру.")

            match = re.search(r'Рисунок\s+(\d+(?:\.\d+)?)', caption)
            if match:
                try:
                    number = int(float(match.group(1)))
                    if number != i:
                        self.issues.append(f"Нарушена последовательность нумерации рисунков. Рисунок с подписью '{caption}' имеет номер {number}, ожидается {i}.")
                except (ValueError, TypeError):
                    self.issues.append(f"Некорректный номер рисунка: '{match.group(1)}' в подписи: '{caption}'")

            has_reference = False
            ref_pattern = r'рисун[окаеи][ак]?\s+' + str(i) + r'\b'

            for paragraph in self.doc.paragraphs:
                if re.search(ref_pattern, paragraph.text, re.IGNORECASE):
                    has_reference = True
                    break

            if not has_reference:
                self.issues.append(f"На рисунок {i} нет ссылки в тексте. На все рисунки должны быть ссылки.")

    def validate_equations(self):
        paragraphs = list(self.doc.paragraphs)
        for i, paragraph in enumerate(paragraphs):
            text = paragraph.text.strip()

            if 'http' in text or 'www.' in text or 'URL:' in text or '[Электронный ресурс]' in text:
                continue

            if re.match(r'^\d+(?:\.\d+)?\.\s+[А-ЯЁа-яёA-Za-z]', text):
                continue

            has_equation_chars = '=' in text
            has_math_operators = any(op in text for op in '+-*/^')
            has_alphabetic_vars = re.search(r'\b[a-zA-Zа-яА-ЯЁё]\s*[=<>]', text) is not None

            if has_equation_chars and (has_math_operators or has_alphabetic_vars):
                prev_index = i - 1
                next_index = i + 1

                if prev_index >= 0 and paragraphs[prev_index].text.strip():
                    if not paragraphs[prev_index].paragraph_format.space_after:
                        self.issues.append(f"Перед формулой '{text[:30]}...' должна быть пустая строка.")

                if next_index < len(paragraphs) and paragraphs[next_index].text.strip():
                    if not paragraph.paragraph_format.space_after:
                        self.issues.append(f"После формулы '{text[:30]}...' должна быть пустая строка.")

                if not re.search(r'\(\d+(?:\.\d+)?\.\d+(?:\.\d+)?\)$', text) and not re.search(r'\(\d+(?:\.\d+)?\)$', text):
                    self.issues.append(f"Формула '{text[:30]}...' должна иметь номер в скобках в конце строки.")

    def validate_references(self):
        required_headings = [
            "СОДЕРЖАНИЕ",
            "ВВЕДЕНИЕ", 
            "ЗАКЛЮЧЕНИЕ", 
            "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ",
            "СПИСОК СОКРАЩЕНИЙ И УСЛОВНЫХ ОБОЗНАЧЕНИЙ",
            "ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ",
        ]
        
        found_headings = {heading: False for heading in required_headings}
        
        for paragraph in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
            text = paragraph.text.strip().upper()
            for heading in required_headings:
                if heading in text:
                    found_headings[heading] = True
        
        for heading in required_headings:
            if not found_headings[heading]:
                self.issues.append(f"Не найден раздел '{heading}'.")
        
        references_paragraphs = []
        in_references_section = False
        
        for paragraph in self.doc.paragraphs:
            if "СПИСОК ИСПОЛЬЗОВАННЫХ ИСТОЧНИКОВ" in paragraph.text.upper():
                in_references_section = True
                continue
            
            if in_references_section and paragraph.text.strip() and any(
                section in paragraph.text.upper() 
                for section in ["ПРИЛОЖЕНИЕ", "СПИСОК СОКРАЩЕНИЙ", "ТЕРМИНЫ И ОПРЕДЕЛЕНИЯ"]
            ):
                in_references_section = False
            
            if in_references_section and paragraph.text.strip():
                references_paragraphs.append(paragraph)
        
        if not references_paragraphs:
            return
        
        for i, paragraph in enumerate(references_paragraphs, 1):
            text = paragraph.text.strip()
            
            has_proper_numbering = False
            
            if hasattr(paragraph, '_p'):
                p_xml = paragraph._p.xml
                if '<w:numPr>' in p_xml and '<w:numId' in p_xml:
                    has_proper_numbering = True
            
            if not has_proper_numbering:
                num_match = re.match(r'^\s*(\d+(?:\.\d+)?)[\.\s]', text)
                if num_match:
                    try:
                        ref_num = int(float(num_match.group(1)))
                        if ref_num != i:
                            self.issues.append(f"Источник '{text[:50]}...' имеет номер {ref_num}, но должен иметь номер {i}.")
                    except (ValueError, TypeError):
                        self.issues.append(f"Некорректный номер источника: '{num_match.group(1)}' в тексте: '{text[:50]}...'")
                else:
                    if not ('http' in text or '[Электронный ресурс]' in text):
                        self.issues.append(f"Источник '{text[:50]}...' должен иметь номер {i}.")
            
            has_citation = False
            citation_pattern = r"\[\s*" + str(i) + r"\s*\]"
            for p in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
                if re.search(citation_pattern, p.text):
                    has_citation = True
                    break
            
            if not has_citation and len(text) > 3:
                self.issues.append(f"На источник #{i} нет ссылки в тексте. На все источники должны быть ссылки.")

    def validate_appendices(self):
        try:
            has_appendices = False
            for paragraph in list(filter(lambda x: x.text != '', self.doc.paragraphs)):
                if paragraph.text.strip().startswith("Приложение "):
                    has_appendices = True
                    if not paragraph.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                        self.issues.append(f"Заголовок '{paragraph.text}' должен быть выровнен по центру.")

                    match = re.match(r'Приложение\s+([А-Я])', paragraph.text.strip())
                    if not match:
                        self.issues.append(f"Неправильный формат обозначения приложения: '{paragraph.text}'. Должно быть 'Приложение' и буква русского алфавита.")
                    else:
                        letter = match.group(1)
                        if letter in 'ЁЗЙОЧЬЫЪ':
                            self.issues.append(f"Недопустимая буква '{letter}' для обозначения приложения.")

            if has_appendices:
                has_reference = False
                for paragraph in self.doc.paragraphs:
                    p_text = paragraph.text.strip()
                    if re.search(r'приложени[еия]', p_text, re.IGNORECASE) and not p_text.startswith("Приложение "):
                        has_reference = True
                        break

                if not has_reference:
                    self.issues.append("На приложения нет ссылок в тексте. На все приложения должны быть ссылки.")
        except Exception as e:
            self.issues.append(f"Ошибка при проверке приложений: {str(e)}")
